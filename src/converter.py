"""
核心转换引擎 — Word (.docx) 与 Markdown (.md) 双向转换。

Word→MD 流程:
  1. python-docx 预处理: 读取样式大纲级别 (w:outlineLvl), 强制写入标准
     "Heading N" 样式名, 保存为临时 .docx
  2. Pandoc 转换: gfm 格式 + 图片提取 + 不折行
  3. Pandoc 不可用时回退到 python-docx 原生解析

MD→Word 流程:
  Pandoc 引擎 (默认) → python-docx 原生引擎 (回退)
"""
import re
import os
import subprocess
import zipfile
from copy import deepcopy
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import markdown as md_lib

from config import load_config


# ── 工具函数 ──────────────────────────────────────────────

def _auto_detect_heading_level(paragraph) -> int | None:
    """
    通过读取样式定义中的大纲级别 (w:outlineLvl) 检测标题层级。

    Word 中每个段落样式可以设置大纲级别：
      - outlineLvl=0  → Level 1  → #  一级标题
      - outlineLvl=1  → Level 2  → ##  二级标题
      - outlineLvl=2  → Level 3  → ###  三级标题
      - outlineLvl=3  → Level 4  → ####  四级标题
      - outlineLvl=4  → Level 5  → #####  五级标题
      - outlineLvl=5  → Level 6  → ######  六级标题
      - 无 outlineLvl  → 正文

    大纲级别存储在样式定义中（w:style/w:pPr/w:outlineLvl），
    因此不同样式名但相同大纲级别 → 同一 MD 标题层级。
    """
    if not paragraph.style:
        return None

    try:
        # 从样式定义的 XML 中读取 w:outlineLvl
        style_elem = paragraph.style._element
        style_pPr = style_elem.find(qn('w:pPr'))
        if style_pPr is not None:
            ol_elem = style_pPr.find(qn('w:outlineLvl'))
            if ol_elem is not None:
                val_str = ol_elem.get(qn('w:val'))
                if val_str is not None:
                    outline_lvl = int(val_str)
                    if 0 <= outline_lvl <= 5:
                        return outline_lvl + 1
    except Exception:
        pass

    return None


def _extract_image_from_docx(doc, image_path: str, output_dir: Path) -> str | None:
    """从 docx 中提取图片到输出目录，返回相对路径。"""
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        with open(image_path, "rb") as img_file:
            img_data = img_file.read()
        # 从 relationship 获取图片文件名
        # image_path 格式类似 word/media/image1.png
        img_name = Path(image_path).name
        dest = output_dir / img_name
        # 避免覆盖
        counter = 1
        while dest.exists():
            stem = Path(image_path).stem
            dest = output_dir / f"{stem}_{counter}{Path(image_path).suffix}"
            counter += 1
        with open(dest, "wb") as f:
            f.write(img_data)
        return str(dest.relative_to(output_dir.parent))
    except Exception:
        return None


def _get_image_rels(doc):
    """获取文档中所有图片关系。"""
    images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images[rel.rId] = rel.target_part.blob
    return images


# ── Word → Markdown ────────────────────────────────────────

def _convert_run_to_md(run) -> str:
    """将 docx 文本运行转换为 Markdown 内联格式。"""
    text = run.text
    if not text:
        return ""

    # 转义 Markdown 特殊字符（不在格式化范围内的）
    # 先保存格式化标记
    is_bold = run.bold
    is_italic = run.italic
    is_underline = run.underline
    is_strike = run.font.strike
    is_code = run.font.name and "consolas" in (run.font.name or "").lower()

    # 对文本中的 Markdown 特殊字符进行转义
    # 在行内代码中不转义（代码内容是字面量）
    if not is_code:
        text = text.replace("\\", "\\\\")
        text = text.replace("*", "\\*")
        text = text.replace("_", "\\_")
        text = text.replace("[", "\\[")
        text = text.replace("]", "\\]")
        text = text.replace("#", "\\#")
        text = text.replace("|", "\\|")
    # 反引号始终转义（避免破坏代码块边界）
    text = text.replace("`", "\\`")

    # 应用内联格式
    if is_bold and is_italic:
        text = f"***{text}***"
    elif is_bold:
        text = f"**{text}**"
    elif is_italic:
        text = f"*{text}*"

    if is_strike:
        text = f"~~{text}~~"

    if is_code:
        text = f"`{text}`"

    return text


def _get_list_info(paragraph, doc=None) -> tuple:
    """获取段落的列表信息，返回 (is_list, level, numId, ordered)。

    检查段落属性、样式定义、以及样式名称来判断列表类型。
    """
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        return False, 0, None, False

    # 1. 直接检查段落中的 numPr
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        ilvl = numPr.find(qn('w:ilvl'))
        numId = numPr.find(qn('w:numId'))
        level = int(ilvl.get(qn('w:val'), 0)) if ilvl is not None else 0
        nid = numId.get(qn('w:val')) if numId is not None else None
        ordered = _is_ordered_numbering(doc, nid) if doc and nid else False
        return True, level, nid, ordered

    # 2. 检查样式名称（常见列表样式）
    if paragraph.style:
        style_name = paragraph.style.name.lower() if paragraph.style.name else ""
        if "list bullet" in style_name or "bullet" in style_name:
            return True, 0, None, False
        if "list number" in style_name or "list num" in style_name:
            return True, 0, None, True

    # 3. 检查样式定义中是否有 numPr
    if paragraph.style and doc:
        style = paragraph.style
        try:
            style_elem = style._element
            style_pPr = style_elem.find(qn('w:pPr'))
            if style_pPr is not None:
                style_numPr = style_pPr.find(qn('w:numPr'))
                if style_numPr is not None:
                    ilvl = style_numPr.find(qn('w:ilvl'))
                    numId = style_numPr.find(qn('w:numId'))
                    level = int(ilvl.get(qn('w:val'), 0)) if ilvl is not None else 0
                    nid = numId.get(qn('w:val')) if numId is not None else None
                    ordered = _is_ordered_numbering(doc, nid) if nid else False
                    return True, level, nid, ordered
        except Exception:
            pass

    return False, 0, None, False


def _is_ordered_numbering(doc, numId: str) -> bool:
    """判断 numbering ID 是否对应有序列表。"""
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return False
        numbering = numbering_part.element
        num = numbering.find(qn(f'w:num[@w:numId="{numId}"]'))
        if num is None:
            return False
        abstractNumId = num.find(qn('w:abstractNumId'))
        if abstractNumId is None:
            return False
        refId = abstractNumId.get(qn('w:val'))
        abstractNum = numbering.find(qn(f'w:abstractNum[@w:abstractNumId="{refId}"]'))
        if abstractNum is None:
            return False
        lvl = abstractNum.find(qn(f'w:lvl[@w:ilvl="0"]'))
        if lvl is None:
            return False
        numFmt = lvl.find(qn('w:numFmt'))
        if numFmt is None:
            return False
        fmt = numFmt.get(qn('w:val'))
        return fmt in ("decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman")
    except Exception:
        return False



def _get_cell_text(cell) -> str:
    """获取表格单元格的纯文本内容。"""
    return cell.text.strip().replace("\n", " ").replace("|", "\\|")


def docx_to_markdown(docx_path: str | Path, config: dict = None,
                     output_path: str | Path = None) -> str:
    """
    将 Word 文档转换为 Markdown 文本。

    优先使用 Pandoc 引擎（效果最佳：表格/列表/图片内联位置精确），
    Pandoc 不可用时回退到 python-docx 原生解析。

    Args:
        docx_path: Word 文档路径
        config: 配置字典（可选，默认从 config.json 加载）
        output_path: 输出 .md 文件路径（Pandoc 路径需要，用于图片路径计算）

    Returns:
        Markdown 格式文本
    """
    if config is None:
        config = load_config()

    docx_path = Path(docx_path)

    # ── Pandoc 路径 ──
    pandoc_exe = _find_pandoc_exe()
    if pandoc_exe and output_path:
        try:
            return _docx_to_md_via_pandoc(docx_path, Path(output_path), config, pandoc_exe)
        except Exception:
            pass  # 回退到原生引擎

    # ── 原生 python-docx 引擎（回退） ──
    return _docx_to_md_native(docx_path, config)


def _find_pandoc_exe() -> str | None:
    """查找 pandoc 可执行文件（复用 pandoc_engine 的查找逻辑）。"""
    try:
        from pandoc_engine import _find_pandoc
        return _find_pandoc()
    except ImportError:
        pass
    import shutil
    return shutil.which("pandoc")


def _docx_to_md_via_pandoc(docx_path: Path, output_path: Path,
                            config: dict, pandoc_exe: str) -> str:
    """
    Pandoc 驱动的 Word→Markdown 转换。

    预处理步骤（lxml + zipfile，不经过 python-docx save）：
      读取 styles.xml，建立 styleId→outlineLvl 映射；
      遍历 document.xml 中每个段落，若其样式的 outlineLvl=0~5，
      则将 w:pStyle 改为标准 "HeadingN"。
      直接操作 XML，完全保留原始 run 结构（<w:br/> 等不被破坏）。

    Pandoc 参数：
      -f docx -t gfm: GitHub Flavored Markdown（表格/列表完美支持）
      --extract-media: 图片提取到指定目录，markdown 中引用路径自动正确
      --wrap=none: 防止中文段落被强制换行打断
    """
    output_dir = output_path.parent
    img_folder = config.get("image_folder", "images")

    # ── 1. 预处理：lxml 直接操作 XML（绕过 python-docx save）──
    temp_docx = output_dir / f"._temp_{docx_path.stem}.docx"
    modified = _fix_headings_via_lxml(docx_path, temp_docx)

    if not modified:
        # 无需修改，直接用原文件
        temp_docx = docx_path

    try:
        # ── 2. 调用 Pandoc ──
        img_dir = output_dir / img_folder
        args = [
            pandoc_exe,
            str(temp_docx),
            "-f", "docx",
            "-t", "gfm",
            "-o", str(output_path),
            f"--extract-media={img_dir}",
            "--wrap=none",
        ]

        result = subprocess.run(
            args,
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            error_msg = result.stderr[:500] if result.stderr else "Unknown Pandoc error"
            raise RuntimeError(f"Pandoc conversion failed: {error_msg}")

        # ── 3. 读取生成的 Markdown ──
        with open(output_path, "r", encoding="utf-8") as f:
            return f.read()

    finally:
        # 清理临时文件
        if modified and temp_docx != docx_path and temp_docx.exists():
            temp_docx.unlink()


# ── lxml + zipfile 实现的 docx 样式修正（不破坏 run 结构）──

# OOXML 命名空间
_WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    """生成带 wml 命名空间的标签名。"""
    return f"{{{_WML_NS}}}{tag}"


def _fix_headings_via_lxml(src_path: Path, dst_path: Path) -> bool:
    """
    通过 lxml + zipfile 直接修改 docx 内部 XML，做三件事：

    1. 标题样式改名
    2. 分节符 → 水平分割线
    3. 手动换行符 → 段落分割

    Returns:
        True 表示有修改，False 表示无需修改（调用方可直接用原文件）。
    """
    # ── 1. 读取原始 zip ──
    with zipfile.ZipFile(src_path, "r") as zf:
        zip_data = {name: zf.read(name) for name in zf.namelist()}

    nsmap = {"w": _WML_NS}
    modified = False

    # ── 2. 样式改名 ──
    styles_root = None
    if "word/styles.xml" in zip_data:
        styles_root = etree.fromstring(zip_data["word/styles.xml"])
        modified |= _rename_heading_styles(styles_root, nsmap)

    # ── 3. <w:sectPr> → 水平分割线 ──
    doc_root = None
    if "word/document.xml" in zip_data:
        doc_root = etree.fromstring(zip_data["word/document.xml"])
        modified |= _convert_section_breaks_to_rules(doc_root, nsmap)

    # ── 4. <w:br/> → 段落分割 ──
    if doc_root is not None:
        modified |= _convert_breaks_to_paragraphs(doc_root, nsmap)

    if not modified:
        return False

    # ── 4. 写入新 zip ──
    with zipfile.ZipFile(dst_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in zip_data.items():
            if name == "word/styles.xml" and styles_root is not None:
                zf.writestr(name, etree.tostring(
                    styles_root, xml_declaration=True, encoding="UTF-8", standalone=True))
            elif name == "word/document.xml" and doc_root is not None:
                zf.writestr(name, etree.tostring(
                    doc_root, xml_declaration=True, encoding="UTF-8", standalone=True))
            else:
                zf.writestr(name, data)

    return True


def _rename_heading_styles(styles_root, nsmap: dict) -> bool:
    """
    将 styles.xml 中 outlineLvl=0~5 但名称非 "heading N" 的样式改名。
    返回 True 表示至少有一个样式被改名。
    """
    modified = False
    for style_elem in styles_root.findall("w:style", nsmap):
        if style_elem.get(_w("type")) != "paragraph":
            continue

        pPr = style_elem.find("w:pPr", nsmap)
        if pPr is None:
            continue
        ol = pPr.find("w:outlineLvl", nsmap)
        if ol is None:
            continue
        val_str = ol.get(_w("val"))
        if val_str is None:
            continue
        try:
            olvl = int(val_str)
        except ValueError:
            continue
        if not (0 <= olvl <= 5):
            continue

        heading_lvl = olvl + 1
        name_elem = style_elem.find("w:name", nsmap)
        if name_elem is None:
            continue
        current_name = name_elem.get(_w("val"), "")
        target_name = f"heading {heading_lvl}"
        if current_name.lower() == target_name:
            continue

        name_elem.set(_w("val"), target_name)
        modified = True

    return modified


def _convert_section_breaks_to_rules(doc_root, nsmap: dict) -> bool:
    """
    将 document.xml 中段落级 <w:sectPr>（分节符）替换为水平线标记。

    OOXML 中分节符出现为 w:p/w:pPr/w:sectPr，Pandoc 可能将其误解析为
    标题。此处移除 w:sectPr 并给段落添加底部边框，Pandoc 会将带底部
    边框的空段落识别为水平分割线（markdown 中即 ---）。

    w:body/w:sectPr 是最后一节的属性，不是分节符，不处理。

    返回 True 表示至少有一个分节符被转换。
    """
    modified = False

    for para in doc_root.findall(".//w:p", nsmap):
        pPr = para.find("w:pPr", nsmap)
        if pPr is None:
            continue
        sectPr = pPr.find("w:sectPr", nsmap)
        if sectPr is None:
            continue

        # 这是段落级分节符
        # 1. 移除 w:sectPr
        pPr.remove(sectPr)

        # 2. 清空段落文本（分节符段落通常是空的）
        for r in para.findall("w:r", nsmap):
            para.remove(r)

        # 3. 添加底部边框（Pandoc 据此识别水平分割线）
        pBdr = pPr.find("w:pBdr", nsmap)
        if pBdr is None:
            pBdr = etree.SubElement(pPr, _w("pBdr"))
        bottom = etree.SubElement(pBdr, _w("bottom"))
        bottom.set(_w("val"), "single")
        bottom.set(_w("sz"), "12")
        bottom.set(_w("space"), "1")
        bottom.set(_w("color"), "auto")

        modified = True

    return modified


def _convert_breaks_to_paragraphs(doc_root, nsmap: dict) -> bool:
    """
    将 document.xml 中所有 <w:br/>（手动换行符）替换为段落分割。

    每个 <w:br/> 所在的段落被拆分为两个独立段落，
    各自继承原段落的 pPr 和对应 run 的 rPr（粗体、斜体等）。
    不包含 <w:br/> 的段落原样保留。

    返回 True 表示至少有一个段落被拆分。
    """
    # 收集所有含 <w:br/> 的段落（避免遍历时修改树）
    paras_with_br = []
    for para in doc_root.findall(".//w:p", nsmap):
        if para.find(".//w:br", nsmap) is not None:
            paras_with_br.append(para)

    if not paras_with_br:
        return False

    for para in paras_with_br:
        new_paras = _split_para_at_breaks(para, nsmap)
        if len(new_paras) <= 1:
            continue
        parent = para.getparent()
        if parent is None:
            continue
        children = list(parent)
        idx = children.index(para)
        parent[idx] = new_paras[0]
        for i, np_elem in enumerate(new_paras[1:], 1):
            parent.insert(idx + i, np_elem)

    return True


def _split_para_at_breaks(para, nsmap: dict) -> list:
    """
    在 <w:br/> 处拆分 w:p，返回新 w:p 元素列表。

    每个段落继承原段落的 pPr（样式等）。含 <w:br/> 的 run
    被拆分：<w:br/> 前的文本留在当前段，后的文本进入下一段，
    各段的 run 都保留原 rPr（粗体/斜体等格式）。
    """
    pPr = para.find("w:pPr", nsmap)

    segments = []     # 最终结果：每个 segment 是一个元素列表
    current = []      # 当前 segment 的元素

    for child in para:
        if child.tag == _w("pPr"):
            continue  # pPr 单独处理

        if child.tag == _w("r"):
            br_elems = child.findall("w:br", nsmap)
            if br_elems:
                # 该 run 含 <w:br/>，拆分
                rPr = child.find("w:rPr", nsmap)
                sub = []  # 当前子段内的 w:t 等
                for rc in child:
                    if rc.tag == _w("br"):
                        # 遇到 br：收束当前子段为一个 run，结束当前 segment
                        if sub:
                            new_r = _make_run_element(rPr, sub)
                            current.append(new_r)
                            sub = []
                        segments.append(current)
                        current = []
                    else:
                        sub.append(deepcopy(rc))
                # 最后一个 br 之后的剩余内容
                if sub:
                    new_r = _make_run_element(rPr, sub)
                    current.append(new_r)
            else:
                current.append(deepcopy(child))
        else:
            current.append(deepcopy(child))

    segments.append(current)

    if len(segments) <= 1:
        return [para]  # 没有拆分

    # 构建新段落
    result = []
    for seg in segments:
        new_para = etree.Element(_w("p"))
        if pPr is not None:
            new_para.append(deepcopy(pPr))
        for elem in seg:
            new_para.append(elem)
        result.append(new_para)

    return result


def _make_run_element(rPr, children: list):
    """创建 w:r 元素，附带 rPr（如有）和子元素列表。"""
    r = etree.Element(_w("r"))
    if rPr is not None:
        r.append(deepcopy(rPr))
    for c in children:
        r.append(c)
    return r


def _docx_to_md_native(docx_path: Path, config: dict) -> str:
    """
    原生 python-docx Word→Markdown 转换（Pandoc 不可用时的回退方案）。

    逐个解析段落/表格/图片，拼接 Markdown 字符串。
    注意：此方案对复杂文档（合并单元格/嵌套列表等）支持有限。
    """
    doc = Document(str(docx_path))
    extract_imgs = config.get("extract_images", True)
    img_folder = config.get("image_folder", "images")

    md_lines = []
    prev_was_empty = False

    # 图片输出目录
    docx_dir = Path(docx_path).parent
    img_out_dir = docx_dir / img_folder

    for para in doc.paragraphs:
        # 1. 检测标题（基于样式的大纲级别）
        heading_level = _auto_detect_heading_level(para)
        if heading_level:
            full_text = "".join(
                _strip_md_escape(run.text) for run in para.runs
            ).strip()
            if full_text:
                md_lines.append(f"{'#' * heading_level} {full_text}")
                md_lines.append("")
                prev_was_empty = True
                continue

        # 2. 检测列表
        is_list, list_level, num_id, is_ordered = _get_list_info(para, doc)
        if is_list and config.get("preserve_lists", True):
            full_text = "".join(
                _strip_md_escape(run.text) for run in para.runs
            ).strip()
            if not full_text:
                continue

            indent = "    " * list_level
            prefix = "1. " if is_ordered else "- "

            md_lines.append(f"{indent}{prefix}{full_text}")
            prev_was_empty = False
            continue

        # 3. 普通段落
        line_parts = []
        for run in para.runs:
            line_parts.append(_convert_run_to_md(run))

        # 处理超链接
        hyperlinks = para._element.findall('.//' + qn('w:hyperlink'))
        for hl in hyperlinks:
            rId = hl.get(qn('r:id'))
            if rId and rId in doc.part.rels:
                rel = doc.part.rels[rId]
                target = rel.target_ref
                line_parts.append(f" (链接: {target})")

        line = "".join(line_parts).strip()

        if line:
            md_lines.append(line)
            prev_was_empty = False
        elif not prev_was_empty:
            md_lines.append("")
            prev_was_empty = True

    # ── 处理表格 ──
    if config.get("preserve_tables", True):
        for table in doc.tables:
            md_lines.append("")
            # 表头
            if table.rows:
                header_cells = [_get_cell_text(cell) for cell in table.rows[0].cells]
                md_lines.append("| " + " | ".join(header_cells) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

                for row in table.rows[1:]:
                    cells = [_get_cell_text(cell) for cell in row.cells]
                    md_lines.append("| " + " | ".join(cells) + " |")

            md_lines.append("")

    # ── 处理图片 ──
    if extract_imgs:
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    img_blob = rel.target_part.blob
                    img_ext = Path(rel.target_part.partname).suffix or ".png"
                    img_name = f"image_{rel_id}{img_ext}"
                    img_out_dir.mkdir(parents=True, exist_ok=True)
                    dest = img_out_dir / img_name
                    counter = 1
                    while dest.exists():
                        img_name = f"image_{rel_id}_{counter}{img_ext}"
                        dest = img_out_dir / img_name
                        counter += 1
                    with open(dest, "wb") as f:
                        f.write(img_blob)
                    md_lines.append(f"![{img_name}]({img_folder}/{img_name})")
                    md_lines.append("")
                except Exception:
                    pass

    return "\n".join(md_lines)


def _strip_md_escape(text: str) -> str:
    """去除 Markdown 转义，恢复纯文本。"""
    text = text.replace("\\\\", "\x00")  # 临时保存 \\
    text = text.replace("\\*", "*")
    text = text.replace("\\_", "_")
    text = text.replace("\\`", "`")
    text = text.replace("\\[", "[")
    text = text.replace("\\]", "]")
    text = text.replace("\\#", "#")
    text = text.replace("\\|", "|")
    text = text.replace("\x00", "\\")
    return text


# ── Markdown → Word ────────────────────────────────────────

def _add_md_paragraph(doc, text: str, style_name: str = None):
    """添加段落到 docx 文档，支持内联格式解析。"""
    para = doc.add_paragraph()

    if style_name:
        try:
            para.style = doc.styles[style_name]
        except KeyError:
            # 如果样式不存在，尝试创建
            pass

    # 解析内联 Markdown: **bold**, *italic*, ~~strike~~, `code`, [link](url)
    # 使用正则逐步解析
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)|'       # bold+italic
        r'(\*\*(.+?)\*\*)|'           # bold
        r'(?<!\*)\*(.+?)\*(?!\*)|'    # italic (single *)
        r'(~~(.+?)~~)|'               # strikethrough
        r'(`(.+?)`)|'                 # inline code
        r'(\[(.+?)\]\((.+?)\))'       # hyperlink
    )

    last_end = 0
    for match in pattern.finditer(text):
        start, end = match.start(), match.end()

        # 添加匹配前的纯文本
        if start > last_end:
            plain = text[last_end:start]
            if plain:
                para.add_run(plain)

        if match.group(1):  # bold+italic
            run = para.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # bold
            run = para.add_run(match.group(4))
            run.bold = True
        elif match.group(5):  # italic
            run = para.add_run(match.group(5))
            run.italic = True
        elif match.group(6):  # strikethrough
            run = para.add_run(match.group(7))
            run.font.strike = True
        elif match.group(8):  # inline code
            run = para.add_run(match.group(9))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            # 添加灰色背景效果 — 通过 shading
            shading = parse_xml(
                f'<w:shd {qn("w:fill")}="F0F0F0" {qn("w:val")}="clear" '
                f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            run._element.get_or_add_rPr().append(shading)
        elif match.group(10):  # hyperlink
            link_text = match.group(11)
            link_url = match.group(12)
            run = para.add_run(link_text)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True

        last_end = end

    # 添加剩余的纯文本
    if last_end < len(text):
        para.add_run(text[last_end:])

    # 如果没有添加任何 run（纯文本无格式），添加整个文本
    if not para.runs:
        para.add_run(text)

    return para


def markdown_to_docx(md_text: str, output_path: str | Path, config: dict = None):
    """
    将 Markdown 文本转换为 Word 文档。

    Args:
        md_text: Markdown 文本内容
        output_path: 输出 .docx 文件路径
        config: 配置字典（可选）
    """
    if config is None:
        config = load_config()

    # ── Pandoc 引擎（默认）──
    try:
        from pandoc_engine import markdown_to_docx_pandoc
        ref_doc = config.get("pandoc_reference_doc", None)
        if markdown_to_docx_pandoc(md_text, str(output_path), ref_doc):
            return  # Pandoc 转换成功
    except ImportError:
        pass  # pandoc_engine 模块不可用，回退到原生

    # ── 原生 python-docx 引擎（Pandoc 不可用时的回退方案）──
    heading_map = config.get("heading_mapping", {}).get("md_to_word", {})

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 确保标题样式存在
    _ensure_heading_styles(doc)

    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    in_quote = False

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith("```"):
            if in_code_block:
                # 结束代码块
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 引用块
        if line.strip().startswith("> "):
            quote_text = line.strip()[2:]
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # 表格
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            # 跳过分隔行
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            table_rows.append(cells)
            i += 1

            # 检查下一行是否还是表格
            if i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                continue
            else:
                # 结束表格，生成 Word 表格
                _add_md_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            prefix = "#" * level
            style_name = heading_map.get(prefix, f"Heading {level}")
            _add_md_paragraph(doc, heading_text, style_name)
            i += 1
            continue

        # 无序列表
        list_match = re.match(r'^(\s*)[\-*+]\s+(.+)$', line)
        if list_match and config.get("preserve_lists", True):
            indent = len(list_match.group(1))
            list_text = list_match.group(2)
            _add_list_item(doc, list_text, indent // 2, ordered=False)
            i += 1
            continue

        # 有序列表
        ordered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ordered_match and config.get("preserve_lists", True):
            indent = len(ordered_match.group(1))
            list_text = ordered_match.group(2)
            _add_list_item(doc, list_text, indent // 4, ordered=True)
            i += 1
            continue

        # 水平线
        if line.strip() in ("---", "***", "___", "* * *"):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            # 添加水平线
            pPr = para._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # 普通段落
        _add_md_paragraph(doc, line)
        i += 1

    # 处理未闭合的代码块
    if in_code_block and code_lines:
        _add_code_block(doc, code_lines)

    # 保存文档
    doc.save(str(output_path))


def _ensure_heading_styles(doc):
    """确保 Word 文档中存在标题样式。"""
    heading_names = [f"Heading {i}" for i in range(1, 7)] + \
                    [f"标题 {i}" for i in range(1, 7)] + \
                    [f"heading {i}" for i in range(1, 7)]
    for name in heading_names:
        try:
            doc.styles[name]
        except KeyError:
            try:
                doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                pass


def _add_code_block(doc, code_lines: list):
    """添加代码块到文档。"""
    code_text = "\n".join(code_lines)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    # 添加灰色背景
    pPr = para._element.get_or_add_pPr()
    shd = parse_xml(
        f'<w:shd {qn("w:fill")}="F5F5F5" {qn("w:val")}="clear" '
        f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    pPr.append(shd)
    run = para.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_list_item(doc, text: str, level: int = 0, ordered: bool = False):
    """添加列表项到文档，支持内联格式。"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    para.paragraph_format.first_line_indent = Inches(-0.25)

    prefix = "" if ordered else "• "
    if ordered:
        prefix = "1. "

    # 先添加前缀作为普通文本
    prefix_run = para.add_run(prefix)

    # 解析并添加格式化内容
    _add_inline_runs(para, text)


def _add_inline_runs(para, text: str):
    """解析内联 Markdown 并添加到段落。"""
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)|'
        r'(\*\*(.+?)\*\*)|'
        r'(?<!\*)\*(.+?)\*(?!\*)|'
        r'(~~(.+?)~~)|'
        r'(`(.+?)`)|'
        r'(\[(.+?)\]\((.+?)\))'
    )

    last_end = 0
    has_match = False
    for match in pattern.finditer(text):
        has_match = True
        start, end = match.start(), match.end()

        if start > last_end:
            para.add_run(text[last_end:start])

        if match.group(1):
            run = para.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):
            run = para.add_run(match.group(4))
            run.bold = True
        elif match.group(5):
            run = para.add_run(match.group(5))
            run.italic = True
        elif match.group(6):
            run = para.add_run(match.group(7))
            run.font.strike = True
        elif match.group(8):
            run = para.add_run(match.group(9))
            run.font.name = "Consolas"
        elif match.group(10):
            run = para.add_run(match.group(11))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True

        last_end = end

    if has_match:
        # 添加最后一个匹配后的剩余文本
        if last_end < len(text):
            para.add_run(text[last_end:])
    else:
        # 无匹配：整个文本作为普通 run
        para.add_run(text)


def _add_md_table(doc, rows: list):
    """添加 Markdown 表格到 Word 文档。"""
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols, style="Table Grid")

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                # 表头加粗
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    doc.add_paragraph()  # 表后空行
